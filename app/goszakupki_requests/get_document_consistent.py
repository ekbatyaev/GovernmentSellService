import os
import sys
import subprocess
import pandas as pd
from PyPDF2 import PdfReader
import docx
import mimetypes
import olefile

# Для Windows DOC
if sys.platform.startswith("win"):
    import win32com.client


def convert_doc_to_docx(doc_path):
    """
    Конвертирует .doc в .docx автоматически (Windows через Word COM,
    Mac/Linux через LibreOffice soffice, если он есть)
    """
    docx_path = os.path.splitext(doc_path)[0] + ".docx"

    if sys.platform.startswith("win"):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, FileFormat=16)  # wdFormatDocumentDefault (.docx)
        doc.Close()
        word.Quit()
    else:
        # Mac / Linux: если установлен LibreOffice
        try:
            subprocess.run([
                "soffice", "--headless", "--convert-to", "docx", doc_path,
                "--outdir", os.path.dirname(doc_path)
            ], check=True)
        except FileNotFoundError:
            raise FileNotFoundError(
                "LibreOffice (soffice) не найден. Для конвертации .doc установите LibreOffice."
            )

    return docx_path


def read_doc_binary(doc_path):
    """Читает текст из старого .doc напрямую через olefile"""
    if not olefile.isOleFile(doc_path):
        raise ValueError(f"Файл не является OLE2 (.doc): {doc_path}")

    ole = olefile.OleFileIO(doc_path)
    text = ""
    if ole.exists('WordDocument'):
        stream = ole.openstream('WordDocument')
        data = stream.read()
        # Простое извлечение ASCII символов
        text = ''.join([chr(b) if 32 <= b <= 126 else '\n' for b in data])
    ole.close()
    return text


def read_file(file_path):
    """Читает текст из одного файла любого популярного формата"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    # TXT
    if ext == '.txt' or mimetypes.guess_type(file_path)[0] == 'text/plain':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    # CSV
    elif ext == '.csv' or mimetypes.guess_type(file_path)[0] == 'text/csv':
        df = pd.read_csv(file_path)
        return df.to_csv(index=False)

    # XLS/XLSX
    elif ext in ['.xls', '.xlsx']:
        df = pd.read_excel(file_path)
        return df.to_csv(index=False)

    # PDF
    elif ext == '.pdf' or mimetypes.guess_type(file_path)[0] == 'application/pdf':
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
        return text.strip()

    # DOCX
    elif ext == '.docx':
        doc = docx.Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs])

    # DOC
    elif ext == '.doc':
        # На Windows пробуем конвертацию через Word COM
        if sys.platform.startswith("win"):
            docx_path = convert_doc_to_docx(file_path)
            doc = docx.Document(docx_path)
            return '\n'.join([p.text for p in doc.paragraphs])
        else:
            # На Mac/Linux читаем напрямую через olefile
            return read_doc_binary(file_path)

    else:
        raise ValueError(f"Формат файла {ext} не поддерживается")


def read_path(path):
    """
    Универсальная функция: если path - файл, возвращает текст;
    если path - папка, возвращает словарь {имя файла: текст}.
    """
    if os.path.isfile(path):
        return read_file(path)
    elif os.path.isdir(path):
        texts = {}
        for root, _, files in os.walk(path):
            for f in files:
                full_path = os.path.join(root, f)
                try:
                    texts[f] = read_file(full_path)
                except Exception as e:
                    print(f"Не удалось прочитать {f}: {e}")
        return texts
    else:
        raise FileNotFoundError(f"Файл или папка не найдены: {path}")


# ======= Пример использования =======
if __name__ == "__main__":
    # Один файл
    text = read_path("tmp/Форма 5.1. Справка о цепочке собственников.xlsx")
    print(text)